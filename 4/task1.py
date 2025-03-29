from docx import Document
from docx.shared import Pt
from docx.shared import Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

head = doc.add_heading('Project Touhou - the best way to kill your PC')
head.style.font.size = Pt(18)
head.style.font.bold = True
head.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph('In this game you need to dodge bullets. It may seem to you quite simple... until you reach the final boss.')
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
fmt = p.paragraph_format
fmt.first_line_indent = Mm(15)
fmt.space_before = Mm(10)
fmt.space_after = Mm(5)

doc.add_paragraph("In the screenshot below you can see different bullet types such as lasers or circles. Of course, touching any them will cause your death... in game.")

doc.add_picture('4/lasers.png', width=Mm(100))

doc.add_paragraph("You may be wondering: \"- and how is that game supposed to kill my computer??\"")
doc.add_paragraph("The answer is obvious: \"- once you get hit by an extra boss you will instantly take your hammer and destroy the monitor of the PC.\"")

doc.save('4/test.docx')